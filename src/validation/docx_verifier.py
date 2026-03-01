from __future__ import annotations

from dataclasses import dataclass
from typing import List, Dict, Optional, Tuple, Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import qn
from docx.table import Table, _Cell
from src.excel.xlsx_reader import read_xlsx_rows

from src.config.config_provider import ConfigProvider
from src.config.constants import (
    WordPlaceholders,
    WordTableDefaults,
    WordLayout,
    XmlTags,
    ConfigKeys,
)
from src.word.placeholder_replacer import get_doc_type_replacements
from src.word.table_handler import _get_docx_path, _find_table_by_header, _row_is_empty


class VerificationError(AssertionError):
    """Raised when document verification fails."""


@dataclass
class CellLocation:
    table_index: int
    row_index: int
    col_index: int

    def __str__(self) -> str:
        return f"table={self.table_index}, row={self.row_index}, col={self.col_index}"


@dataclass
class CellDiff:
    location: CellLocation
    expected: str
    actual: str
    reason: str


def _load_document(path: str) -> Document:
    """Load a .docx with the same resolution rules as the table handler."""
    return Document(_get_docx_path(path))


def _iter_all_paragraphs(doc: Document):
    """Yield all paragraphs in body, tables, headers and footers."""
    # Body paragraphs
    for p in doc.paragraphs:
        yield ("body", None, None, None, p)

    # Body tables
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    yield ("table", ti, ri, ci, p)

    # Headers & footers
    for si, section in enumerate(doc.sections):
        header = section.header
        for p in header.paragraphs:
            yield (f"header[{si}]", None, None, None, p)
        for ti, table in enumerate(header.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    for p in cell.paragraphs:
                        yield (f"header[{si}].table", ti, ri, ci, p)

        footer = section.footer
        for p in footer.paragraphs:
            yield (f"footer[{si}]", None, None, None, p)
        for ti, table in enumerate(footer.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    for p in cell.paragraphs:
                        yield (f"footer[{si}].table", ti, ri, ci, p)


def _normalize_text(text: str) -> str:
    """
    Semantic text normalisation:
    - Strip leading/trailing whitespace
    - Collapse internal whitespace (incl. non‑breaking) to single spaces
    """
    if text is None:
        return ""
    # Normalise various whitespace characters that Word commonly uses
    # (nbsp, tabs, newlines) into single spaces, then collapse.
    import re

    cleaned = (
        text.replace("\u00A0", " ")  # non-breaking space
        .replace("\t", " ")
        .replace("\r", " ")
        .replace("\n", " ")
    )
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def _cell_logical_text(cell: _Cell) -> str:
    """
    Build a logical text representation for a cell:
    - Preserve paragraph boundaries with '\n'
    - Concatenate runs within a paragraph
    - Then apply semantic normalisation
    """
    parts: List[str] = []
    for p in cell.paragraphs:
        # Paragraph.text already concatenates runs, but we prefer explicit control
        p_text = "".join(r.text for r in p.runs) or p.text
        parts.append(p_text or "")
    joined = "\n".join(parts)
    return _normalize_text(joined)


def _has_merged_cells(table: Table) -> bool:
    """
    Detect merged cells via gridSpan (horizontal) or vMerge (vertical).
    Any presence means the table is structurally unsafe for naïve row/column comparison.
    """
    for row in table._tbl.tr_lst:
        for tc in row.tc_lst:
            tc_pr = tc.tcPr
            if tc_pr is None:
                continue
            grid_span = tc_pr.find(qn(XmlTags.CELL_WIDTH.replace("tcW", "gridSpan")))
            v_merge = tc_pr.find(qn("w:vMerge"))
            if grid_span is not None or v_merge is not None:
                return True
    return False


def _assert_no_merged_cells(table: Table, table_label: str) -> None:
    if _has_merged_cells(table):
        raise VerificationError(
            f"Merged cells detected in {table_label}. "
            "Merged cells are not supported by the normalization pipeline and would break row integrity."
        )


def _collect_table_matrix(table: Table, skip_header: bool) -> List[List[str]]:
    """
    Collect a 2D logical text matrix from a python-docx Table.
    Empty rows (all cells empty/whitespace) are treated as data rows with empty strings.
    """
    rows = table.rows[1:] if (skip_header and len(table.rows) > 0) else table.rows
    matrix: List[List[str]] = []
    for row in rows:
        matrix.append([_cell_logical_text(cell) for cell in row.cells])
    return matrix


def _collect_excel_matrix(path: str, sheet_name: Optional[str] = None, skip_header: bool = True) -> List[List[str]]:
    """
    Collect a 2D logical text matrix from an Excel worksheet.
    Preserves empty columns by ensuring all rows have the same length.
    """
    rows = read_xlsx_rows(path, sheet_name=sheet_name)

    if not rows:
        return []

    # Determine the maximum column count across all rows
    # This ensures we don't lose empty columns
    max_cols = max(len(row) for row in rows) if rows else 0

    # Normalize all rows to have the same column count
    # Fill missing cells with empty strings
    normalized_rows = []
    for row in rows:
        normalized_row = list(row) + [''] * (max_cols - len(row))
        normalized_rows.append([_normalize_text(str(cell) if cell is not None else '') for cell in normalized_row])

    if skip_header and normalized_rows:
        return normalized_rows[1:]

    return normalized_rows


def _table_matches_headers(table: Table, expected_headers: List[str]) -> bool:
    """Check if table's first row matches expected headers."""
    if not table.rows:
        return False
    header_cells = [_normalize_text(cell.text) for cell in table.rows[0].cells]
    expected_normalized = [_normalize_text(h) for h in expected_headers]
    return header_cells == expected_normalized


def _get_placeholder_replacements() -> Dict[str, str]:
    """Get all placeholder replacements from config including doc_type overrides."""
    config = ConfigProvider.load_config_json()

    replacements = {
        WordPlaceholders.DOC_TYPE: config.get(
            ConfigKeys.DOC_TYPE, config.get(ConfigKeys.LEGACY_KEYS["DOC_TYPE"], "")
        ),
        WordPlaceholders.DOC_TYPE_STx: config.get(
            ConfigKeys.DOC_STX, config.get(ConfigKeys.LEGACY_KEYS["DOC_TYPE_STX"], "")
        ),
        WordPlaceholders.DOC_RECORD: config.get(
            ConfigKeys.DOC_RECORD, config.get(ConfigKeys.LEGACY_KEYS["DOC_RECORD"], "")
        ),
        WordPlaceholders.DOC_STD: config.get(
            ConfigKeys.DOC_STD, config.get(ConfigKeys.LEGACY_KEYS["DOC_STD"], "")
        ),
        WordPlaceholders.STD_NAME: config.get(
            ConfigKeys.STD_NAME, config.get(ConfigKeys.LEGACY_KEYS["STD_NAME"], "")
        ),
        WordPlaceholders.PLAN_NUMBER: config.get(
            ConfigKeys.TEST_PLAN, config.get(ConfigKeys.LEGACY_KEYS["PLAN_NUMBER"], "")
        ),
        WordPlaceholders.PREPARED_BY: config.get(
            ConfigKeys.PREPARED_BY, config.get(ConfigKeys.LEGACY_KEYS["PREPARED_BY"], "")
        ),
        WordPlaceholders.TEST_PROTOCOL: config.get(
            ConfigKeys.TEST_PROTOCOL, config.get(ConfigKeys.LEGACY_KEYS["TEST_PROTOCOL"], "")
        ),
        WordPlaceholders.FOOTER: config.get(
            ConfigKeys.FOOTER, config.get(ConfigKeys.LEGACY_KEYS["FOOTER"], "")
        ),
    }

    # Apply doc_type-based overrides
    doc_type_from_config = config.get(ConfigKeys.DOC_TYPE) or config.get(
        ConfigKeys.LEGACY_KEYS["DOC_TYPE"]
    )
    doc_type_replacements = get_doc_type_replacements(doc_type_from_config)
    if doc_type_replacements:
        replacements.update(doc_type_replacements)

    return replacements


def _is_valid_placeholder_replacement(template_text: str, normalized_text: str) -> bool:
    """
    Check if difference between template and normalized text is due to valid placeholder replacement.
    Returns True if the normalized text matches what we'd expect after replacing all placeholders.
    """
    replacements = _get_placeholder_replacements()

    # Try replacing all placeholders in template text
    expected = template_text
    for placeholder, value in replacements.items():
        if value:
            expected = expected.replace(placeholder, value)

    return _normalize_text(expected) == _normalize_text(normalized_text)


def validate_table_content_integrity(
        exported_std_path: str,
        normalized_protocol_path: str,
        expected_target_headers: Optional[List[str]] = None,
) -> None:
    """
    PROVE that every logical data cell from the source STD exists in the output document.

    Validation logic:
    - Source: first table in exported STD, data rows = all rows after header.
    - Target: table in normalized protocol identified by header detection.
    - Target data rows = all non-empty rows after header.
    - Row and column counts must match.
    - Each corresponding logical cell (semantic text) must match.
    - Detect merged cells and fail explicitly if present.
    """
    source_is_excel = exported_std_path.lower().endswith(".xlsx")
    exported_doc = None if source_is_excel else _load_document(exported_std_path)
    normalized_doc = _load_document(normalized_protocol_path)

    if source_is_excel:
        src_matrix = _collect_excel_matrix(exported_std_path, skip_header=True)
    else:
        if not exported_doc.tables:
            raise VerificationError("Source STD document contains no tables.")

        src_table = exported_doc.tables[0]
        _assert_no_merged_cells(src_table, "source STD table[0]")
        src_matrix = _collect_table_matrix(src_table, skip_header=True)

    if expected_target_headers is None:
        expected_target_headers = WordTableDefaults.DEFAULT_TARGET_HEADERS

    try:
        target_table = _find_table_by_header(normalized_doc, expected_headers=expected_target_headers)
    except Exception as exc:
        raise VerificationError(
            f"Target table with headers {expected_target_headers} not found in normalized protocol: {exc}"
        ) from exc

    _assert_no_merged_cells(target_table, "normalized protocol target table")

    # In the target table, ignore purely empty rows (template placeholders that were not used)
    # but preserve order.
    target_rows_after_header = list(target_table.rows)[1:] if len(target_table.rows) > 1 else []
    target_data_rows: List[List[str]] = []
    for row in target_rows_after_header:
        if _row_is_empty(row):
            continue
        target_data_rows.append([_cell_logical_text(cell) for cell in row.cells])

    # Row count check
    if len(src_matrix) != len(target_data_rows):
        raise VerificationError(
            f"Row count mismatch between source STD and normalized protocol target table. "
            f"Expected {len(src_matrix)} data rows from STD, "
            f"but found {len(target_data_rows)} non-empty data rows in target table."
        )

    # Column count & per-cell checks with rich diagnostics
    for row_idx, (src_row, dst_row) in enumerate(zip(src_matrix, target_data_rows), start=1):
        if len(src_row) != len(dst_row):
            raise VerificationError(
                f"Column count mismatch at data row {row_idx}. "
                f"Source has {len(src_row)} columns, target has {len(dst_row)} columns."
            )
        for col_idx, (src_cell, dst_cell) in enumerate(zip(src_row, dst_row), start=0):
            if src_cell != dst_cell:
                location = CellLocation(
                    table_index=0,  # source is fixed; target table is identified by header
                    row_index=row_idx,
                    col_index=col_idx,
                )
                raise VerificationError(
                    f"Content mismatch at {location}. "
                    f"Expected (from STD): '{src_cell}' | Actual (in protocol): '{dst_cell}'."
                )


def validate_structural_correctness(
        normalized_protocol_path: str,
        expected_target_headers: Optional[List[str]] = None,
        expected_column_count: Optional[int] = None,
) -> None:
    """
    Validate structural correctness of the target table:
    - Table is located via header detection.
    - Column count matches expected (defaults to config widths length).
    - Ensures there is at least one data row present (insertion took place).
    """
    doc = _load_document(normalized_protocol_path)

    if expected_target_headers is None:
        expected_target_headers = WordTableDefaults.DEFAULT_TARGET_HEADERS

    try:
        table = _find_table_by_header(doc, expected_headers=expected_target_headers)
    except Exception as exc:
        raise VerificationError(
            f"Target table with headers {expected_target_headers} not found in normalized protocol: {exc}"
        ) from exc

    if not table.rows:
        raise VerificationError("Target table has no rows at all (missing header row).")

    header_cells = table.rows[0].cells
    if expected_column_count is None:
        expected_column_count = len(WordTableDefaults.DEFAULT_COLUMN_WIDTHS_CM)

    if len(header_cells) != expected_column_count:
        raise VerificationError(
            f"Target table column count mismatch. "
            f"Expected {expected_column_count} columns, found {len(header_cells)}."
        )

    # Ensure that we actually have inserted data rows (either into placeholder rows or appended)
    data_rows = list(table.rows)[1:] if len(table.rows) > 1 else []
    non_empty_data_rows = [r for r in data_rows if not _row_is_empty(r)]
    if not non_empty_data_rows:
        raise VerificationError(
            "Target table contains no non-empty data rows after the header. "
            "Row insertion into placeholder/appended rows appears to have failed."
        )


def validate_formatting(
        normalized_protocol_path: str,
        expected_column_widths_cm: Optional[List[float]] = None,
        expected_target_headers: Optional[List[str]] = None,
) -> None:
    """
    Validate formatting normalization:
    - All sections are landscape.
    - All tables are set to AutoFit to Window (tblW type=pct, w=5000).
    - Target table has fixed column widths matching the configuration.
    - Paragraph spacing (before, after) is Normalized (0pt before, 3pt after).
    - Second column paragraphs:
        * Have 'Normal' style.
        * Have no numbering (w:numPr / w:outlineLvl).
    """
    from docx.shared import Pt

    doc = _load_document(normalized_protocol_path)

    # ---- Landscape orientation on all sections ----
    for idx, section in enumerate(doc.sections):
        if section.orientation != WD_ORIENT.LANDSCAPE:
            raise VerificationError(f"Section {idx} is not in landscape orientation.")
        if section.page_width < section.page_height:
            raise VerificationError(
                f"Section {idx} has portrait page dimensions (width < height) despite LANDSCAPE orientation."
            )

    # ---- Target table column widths ----
    if expected_target_headers is None:
        expected_target_headers = WordTableDefaults.DEFAULT_TARGET_HEADERS
    if expected_column_widths_cm is None:
        expected_column_widths_cm = WordTableDefaults.DEFAULT_COLUMN_WIDTHS_CM

    try:
        target_table = _find_table_by_header(doc, expected_headers=expected_target_headers)
    except Exception as exc:
        raise VerificationError(
            f"Target table with headers {expected_target_headers} not found when validating column widths: {exc}"
        ) from exc

    if not target_table.rows:
        raise VerificationError("Target table has no rows when validating column widths.")

    if len(target_table.rows[0].cells) != len(expected_column_widths_cm):
        raise VerificationError(
            "Target table column count does not match configured widths length: "
            f"{len(target_table.rows[0].cells)} vs {len(expected_column_widths_cm)}."
        )

    # Validate paragraph spacing across the entire document
    expected_before = Pt(WordTableDefaults.DEFAULT_PARAGRAPH_SPACING_BEFORE_PT)
    expected_after = Pt(WordTableDefaults.DEFAULT_PARAGRAPH_SPACING_AFTER_PT)

    def _check_spacing(paragraph, context: str) -> None:
        pf = paragraph.paragraph_format
        # python-docx may represent these as Length or None; normalise via .pt
        before = getattr(pf.space_before, "pt", 0 if pf.space_before is None else pf.space_before)
        after = getattr(pf.space_after, "pt", 0 if pf.space_after is None else pf.space_after)
        if abs(before - expected_before.pt) > 0.01 or abs(after - expected_after.pt) > 0.01:
            raise VerificationError(
                f"Paragraph spacing mismatch in {context}. "
                f"Expected before={expected_before.pt}pt, after={expected_after.pt}pt; "
                f"got before={before}pt, after={after}pt."
            )

    # Body paragraphs
    for idx, p in enumerate(doc.paragraphs):
        _check_spacing(p, f"body paragraph {idx}")

    # Table paragraphs
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    _check_spacing(p, f"table {ti} row {ri} col {ci} paragraph {pi}")

    # ---- Second column formatting in the target table ----
    for ri, row in enumerate(target_table.rows[1:], start=1):  # skip header
        if len(row.cells) < 2:
            continue
        cell = row.cells[1]
        for pi, p in enumerate(cell.paragraphs):
            # Style name check (semantic)
            style_name = getattr(p.style, "name", None)
            if style_name is None or style_name.lower() != "normal":
                raise VerificationError(
                    f"Second column paragraph at row {ri}, paragraph {pi} "
                    f"has style '{style_name}', expected 'Normal'."
                )

            # Numbering check via underlying XML
            p_elm = p._element
            pPr = p_elm.find(qn(XmlTags.PARAGRAPH_PROPERTIES))
            if pPr is not None:
                numPr = pPr.find(qn(XmlTags.NUMBERING_PROPERTIES))
                outlineLvl = pPr.find(qn(XmlTags.OUTLINE_LEVEL))
                if numPr is not None or outlineLvl is not None:
                    raise VerificationError(
                        f"Second column paragraph at row {ri}, paragraph {pi} still has numbering metadata "
                        f"(w:numPr or w:outlineLvl present)."
                    )


def detect_unresolved_placeholders(doc: Document) -> Dict[str, List[str]]:
    """
    Scan the entire document (body, tables, headers, footers) and return
    a mapping: placeholder_token -> list of human-readable locations where it appears.
    """
    placeholders = [
        WordPlaceholders.DOC_TYPE,
        WordPlaceholders.DOC_TYPE_STx,
        WordPlaceholders.DOC_RECORD,
        WordPlaceholders.DOC_STD,
        WordPlaceholders.STD_NAME,
        WordPlaceholders.PLAN_NUMBER,
        WordPlaceholders.PREPARED_BY,
        WordPlaceholders.TEST_PROTOCOL,
        WordPlaceholders.FOOTER,
    ]

    hits: Dict[str, List[str]] = {ph: [] for ph in placeholders}

    for scope, ti, ri, ci, p in _iter_all_paragraphs(doc):
        text = p.text or ""
        for ph in placeholders:
            if ph in text:
                location = f"{scope}"
                if ti is not None:
                    location += f", table={ti}"
                if ri is not None and ci is not None:
                    location += f", row={ri}, col={ci}"
                hits[ph].append(location)

    # Remove placeholders that were not found
    return {ph: locs for ph, locs in hits.items() if locs}


def validate_placeholder_replacement(normalized_protocol_path: str) -> None:
    """
    Validate that:
    - All known placeholders are fully replaced in body, tables, headers, and footers.
    - Replacement values exactly match configuration (including doc_type overrides).
    """
    doc = _load_document(normalized_protocol_path)

    unresolved = detect_unresolved_placeholders(doc)
    if unresolved:
        details = "; ".join(
            f"{ph} -> {locations}" for ph, locations in unresolved.items()
        )
        raise VerificationError(
            f"Unresolved placeholders remain in the normalized protocol: {details}"
        )

    # Compute expected replacement values using the same logic as the production replacer
    replacements = _get_placeholder_replacements()

    # Gather full text of the document for existence checks
    all_text_fragments: List[str] = []
    for _, _, _, _, p in _iter_all_paragraphs(doc):
        if p.text:
            all_text_fragments.append(p.text)
    full_text = "\n".join(all_text_fragments)

    for placeholder, expected_value in replacements.items():
        if not expected_value:
            continue  # empty config is allowed; nothing to assert
        if expected_value not in full_text:
            raise VerificationError(
                f"Expected replacement value for {placeholder} not found in document text. "
                f"Expected: '{expected_value}'."
            )


def validate_template_preservation(
        template_protocol_path: str,
        normalized_protocol_path: str,
        expected_target_headers: Optional[List[str]] = None,
) -> None:
    """
    Verify that all non-target-table content from the template
    is preserved in the normalized output.

    This ensures:
    - Table count remains the same
    - Non-target tables are identical (allowing only placeholder replacements)
    - Table structure (row/column counts) is preserved
    """
    template_doc = _load_document(template_protocol_path)
    normalized_doc = _load_document(normalized_protocol_path)

    if expected_target_headers is None:
        expected_target_headers = WordTableDefaults.DEFAULT_TARGET_HEADERS

    # Find target table indices
    template_target_idx = None
    for idx, table in enumerate(template_doc.tables):
        if _table_matches_headers(table, expected_target_headers):
            template_target_idx = idx
            break

    normalized_target_idx = None
    for idx, table in enumerate(normalized_doc.tables):
        if _table_matches_headers(table, expected_target_headers):
            normalized_target_idx = idx
            break

    if template_target_idx is None:
        raise VerificationError(
            f"Target table with headers {expected_target_headers} not found in template."
        )

    if normalized_target_idx is None:
        raise VerificationError(
            f"Target table with headers {expected_target_headers} not found in normalized document."
        )

    # Verify table count (should be same)
    if len(template_doc.tables) != len(normalized_doc.tables):
        raise VerificationError(
            f"Table count mismatch. Template has {len(template_doc.tables)} tables, "
            f"normalized has {len(normalized_doc.tables)} tables."
        )

    # Verify non-target tables are identical (allowing placeholder replacements)
    for idx, (tmpl_table, norm_table) in enumerate(zip(template_doc.tables, normalized_doc.tables)):
        if idx == template_target_idx:
            continue  # Skip target table (it's intentionally modified)

        # Compare structure
        if len(tmpl_table.rows) != len(norm_table.rows):
            raise VerificationError(
                f"Non-target table {idx} row count changed. "
                f"Template: {len(tmpl_table.rows)}, Normalized: {len(norm_table.rows)}"
            )

        # Compare content
        for ri, (tmpl_row, norm_row) in enumerate(zip(tmpl_table.rows, norm_table.rows)):
            if len(tmpl_row.cells) != len(norm_row.cells):
                raise VerificationError(
                    f"Non-target table {idx} row {ri} column count changed. "
                    f"Template: {len(tmpl_row.cells)}, Normalized: {len(norm_row.cells)}"
                )
            for ci, (tmpl_cell, norm_cell) in enumerate(zip(tmpl_row.cells, norm_row.cells)):
                tmpl_text = _cell_logical_text(tmpl_cell)
                norm_text = _cell_logical_text(norm_cell)

                # If texts differ, check if it's a valid placeholder replacement
                if tmpl_text != norm_text:
                    if not _is_valid_placeholder_replacement(tmpl_text, norm_text):
                        raise VerificationError(
                            f"Non-target table {idx} cell (row={ri}, col={ci}) content changed unexpectedly. "
                            f"Template: '{tmpl_text}' -> Normalized: '{norm_text}'"
                        )


def validate_body_paragraphs_preserved(
        template_protocol_path: str,
        normalized_protocol_path: str,
) -> None:
    """
    Verify body paragraphs (outside tables) are preserved.

    Allows only placeholder replacements as valid differences.
    """
    template_doc = _load_document(template_protocol_path)
    normalized_doc = _load_document(normalized_protocol_path)

    if len(template_doc.paragraphs) != len(normalized_doc.paragraphs):
        raise VerificationError(
            f"Body paragraph count changed. "
            f"Template: {len(template_doc.paragraphs)}, "
            f"Normalized: {len(normalized_doc.paragraphs)}"
        )

    for idx, (tmpl_p, norm_p) in enumerate(zip(template_doc.paragraphs, normalized_doc.paragraphs)):
        tmpl_text = _normalize_text(tmpl_p.text)
        norm_text = _normalize_text(norm_p.text)

        if tmpl_text != norm_text:
            if not _is_valid_placeholder_replacement(tmpl_p.text, norm_p.text):
                raise VerificationError(
                    f"Body paragraph {idx} changed unexpectedly. "
                    f"Template: '{tmpl_text}' -> Normalized: '{norm_text}'"
                )


def verify_normalized_protocol(
        exported_std_path: str,
        template_protocol_path: str,
        normalized_protocol_path: str,
) -> None:
    """
    High-level, reusable verification entry point for CI-grade validation.

    This orchestrates:
    - Content integrity checks between the exported STD and normalized protocol.
    - Structural correctness of the target table.
    - Formatting normalization across the document.
    - Placeholder replacement correctness.
    - Template preservation (non-target content unchanged).

    Any deviation raises VerificationError with a precise, human-readable root cause.
    """
    # 1. Content integrity & structural correctness (tables & rows)
    validate_table_content_integrity(
        exported_std_path=exported_std_path,
        normalized_protocol_path=normalized_protocol_path,
    )
    validate_structural_correctness(normalized_protocol_path=normalized_protocol_path)

    # 2. Formatting normalization
    validate_formatting(normalized_protocol_path=normalized_protocol_path)

    # 3. Placeholders
    validate_placeholder_replacement(normalized_protocol_path=normalized_protocol_path)

    # 4. Template preservation - verify non-target tables remain unchanged
    validate_template_preservation(
        template_protocol_path=template_protocol_path,
        normalized_protocol_path=normalized_protocol_path,
    )

    # 5. Body paragraphs preservation - verify paragraphs outside tables remain unchanged
    validate_body_paragraphs_preserved(
        template_protocol_path=template_protocol_path,
        normalized_protocol_path=normalized_protocol_path,
    )
