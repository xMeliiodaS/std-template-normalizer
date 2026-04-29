from copy import deepcopy
import os
import re
from docx import Document
from docx.shared import Cm, Pt
from docx.table import Table
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement, qn
from src.excel.xlsx_reader import read_xlsx_rows


def _ensure_docx_extension(path: str) -> str:
    """Ensure the path has .docx extension. Returns the path with .docx if it doesn't have it."""
    if path and not path.endswith('.docx'):
        return path + '.docx'
    return path


def _get_docx_path(path: str) -> str:
    """Get the docx path, checking if .docx extension needs to be added for opening."""
    if path.endswith('.docx'):
        return path
    # Check if file exists with .docx extension
    if os.path.exists(path + '.docx'):
        return path + '.docx'
    # If file exists without extension, try to open it (python-docx can handle it)
    if os.path.exists(path):
        return path
    # If neither exists, try with .docx extension (might be creating new file)
    return path + '.docx'


from pathlib import Path

def set_landscape_for_all_sections(docx_path: str, output_path: str = None):
    """
    Force all sections to Landscape without toggling.
    - Sets orientation to LANDSCAPE.
    - Swaps width/height only if the current page is Portrait (width < height).
    """
    document = Document(_get_docx_path(docx_path))

    for section in document.sections:
        # Always set orientation
        section.orientation = WD_ORIENT.LANDSCAPE

        # Only swap if the page is currently portrait
        if section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width

    save_path = _ensure_docx_extension(output_path or docx_path)

    # Ensure the target folder exists before saving
    Path(save_path).parent.mkdir(parents=True, exist_ok=True)

    document.save(save_path)

def set_tables_autofit_to_window(docx_path: str, output_path: str = None, clear_column_widths: bool = True):
    """
    Applies Word's 'Layout -> AutoFit -> AutoFit to Window' to all tables in a .docx:
      - Set table preferred width to 100% (pct=5000, meaning 100%).
      - Remove fixed table layout (w:tblLayout w:type='fixed') to enable AutoFit behavior.
      - Optionally remove cell-level fixed widths so Word can reflow columns.

    Args:
        docx_path: str - path to input .docx
        output_path: str | None - path to save; overwrites input if None
        clear_column_widths: bool - remove <w:tcW> from cells to allow true autofit
    """
    doc = Document(_get_docx_path(docx_path))

    for table in doc.tables:
        tbl = table._tbl  # <w:tbl>
        tblPr = tbl.tblPr
        if tblPr is None:
            # Ensure <w:tblPr> exists
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # ---- Preferred width: 100% (AutoFit to Window) ----
        # Look for existing <w:tblW>; create if missing
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        # Set as percentage (fiftieths of a percent): 100% -> 5000
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')

        # ---- Remove fixed table layout to allow AutoFit ----
        # If <w:tblLayout w:type="fixed"> exists, remove it
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is not None:
            tblPr.remove(tblLayout)

        # ---- Optional: Clear per-cell fixed widths ----
        if clear_column_widths:
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is not None:
                        tcPr.remove(tcW)

    save_path = _ensure_docx_extension(output_path or docx_path)
    Path(save_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)


def _row_is_empty(row) -> bool:
    """Return True if all cell texts in the row are empty/whitespace."""
    for cell in row.cells:
        if cell.text and cell.text.strip():
            return False
    return True


def _table_header_matches(table: Table, expected_headers: list[str], case_insensitive: bool = True) -> bool:
    """
    Check whether the table's first row cells match expected header labels.
    Supports partial matching when len(expected_headers) <= number of cells.
    """
    if len(table.rows) == 0:
        return False
    header_cells = table.rows[0].cells
    if len(header_cells) < len(expected_headers):
        return False

    for i, exp in enumerate(expected_headers):
        actual = header_cells[i].text.strip() if i < len(header_cells) else ""
        if case_insensitive:
            if exp.strip().lower() not in actual.lower():
                return False
        else:
            if exp.strip() not in actual:
                return False
    return True


def _find_table_by_header(dst_doc: Document, expected_headers: list[str]) -> Table:
    """
    Return the first table whose header row matches expected_headers.
    Example expected_headers: ["ID"] or ["ID", "Name", "Description"].
    """
    for table in dst_doc.tables:
        if _table_header_matches(table, expected_headers):
            return table
    raise ValueError(f"No table found with header(s): {expected_headers}")


def copy_table_rows_excluding_header_into_table_with_id(
        src_docx_path: str,
        dst_docx_path: str,
        output_path: str = None,
        src_table_index: int = 0,
        expected_target_headers=None
):
    """
    Copy all rows except the header from a source table and paste into the target table
    identified by its header row (e.g., header contains 'ID').
    Insertion starts at the first empty row below the header; if none, rows are appended.

    Args:
        src_docx_path: path to source .docx (table to copy from)
        dst_docx_path: path to target .docx (table to paste into)
        output_path: where to save updated target; overwrites dst_docx_path if None
        src_table_index: 0-based index of source table
        expected_target_headers: list of header labels to identify target table
    """
    # Load docs
    if expected_target_headers is None:
        expected_target_headers = ["ID"]
    src = Document(_get_docx_path(src_docx_path))
    dst = Document(_get_docx_path(dst_docx_path))

    # Validate source table
    if src_table_index < 0 or src_table_index >= len(src.tables):
        raise IndexError(f"Source table index {src_table_index} out of range. Source has {len(src.tables)} tables.")

    src_table = src.tables[src_table_index]
    src_tbl_elm = src_table._tbl

    # Gather source rows excluding header (skip first <w:tr>)
    src_rows = src_tbl_elm.findall(qn('w:tr'))
    if len(src_rows) < 2:
        raise ValueError("Source table must have at least 2 rows (header + data).")
    rows_to_copy = [deepcopy(tr) for tr in src_rows[1:]]

    # Defensive: clear repeating header flags on copied rows
    for tr in rows_to_copy:
        trPr = tr.find(qn('w:trPr'))
        if trPr is not None:
            hdr = trPr.find(qn('w:tblHeader'))
            if hdr is not None:
                trPr.remove(hdr)

    # Find target table by its header row (e.g., header first cell contains 'ID')
    target_table = _find_table_by_header(dst, expected_headers=expected_target_headers)
    target_tbl_elm = target_table._tbl

    # Identify first empty data row (below header)
    empty_row_obj = None
    data_rows = list(target_table.rows)[1:] if len(target_table.rows) > 1 else []
    for r in data_rows:
        if _row_is_empty(r):
            empty_row_obj = r
            break

    if empty_row_obj is not None:
        # Replace the empty row with the first copied row, then insert remaining
        empty_tr = empty_row_obj._tr
        empty_tr.addnext(rows_to_copy[0])
        target_tbl_elm.remove(empty_tr)
        cursor = rows_to_copy[0]
        for tr in rows_to_copy[1:]:
            cursor.addnext(tr)
            cursor = tr
    else:
        # No empty row found → append at end
        for tr in rows_to_copy:
            target_tbl_elm.append(tr)

    set_normal_style_in_second_column(target_table, skip_header=True)
    remove_numbering_in_second_column(target_table, skip_header=True)

    save_path = _ensure_docx_extension(output_path or dst_docx_path)
    Path(save_path).parent.mkdir(parents=True, exist_ok=True)
    dst.save(save_path)


def copy_excel_rows_excluding_header_into_table_with_id(
        src_excel_path: str,
        dst_docx_path: str,
        output_path: str = None,
        sheet_name: str = None,
        expected_target_headers=None,
):
    """
    Copy all rows except the header from an Excel sheet into the Word target table
    identified by its header row (e.g., first cell is 'ID').
    """
    if expected_target_headers is None:
        expected_target_headers = ["ID"]

    source_rows = read_xlsx_rows(src_excel_path, sheet_name=sheet_name)
    if len(source_rows) < 2:
        raise ValueError("Source Excel sheet must have at least 2 rows (header + data).")

    data_rows = source_rows[1:]

    dst = Document(_get_docx_path(dst_docx_path))
    target_table = _find_table_by_header(dst, expected_headers=expected_target_headers)

    # Reuse first empty row if template provides one, otherwise append rows.
    template_row = None
    data_rows_in_target = list(target_table.rows)[1:] if len(target_table.rows) > 1 else []
    for row in data_rows_in_target:
        if _row_is_empty(row):
            template_row = row
            break

    if template_row is None:
        target_table.add_row()
        template_row = target_table.rows[-1]

    template_row_xml = template_row._tr
    target_tbl_elm = target_table._tbl

    for row_idx, source_row in enumerate(data_rows):
        if row_idx == 0:
            current_row = template_row
        else:
            cloned_tr = deepcopy(template_row_xml)
            target_tbl_elm.append(cloned_tr)
            current_row = target_table.rows[-1]

        for col_idx, value in enumerate(source_row):
            if col_idx >= len(current_row.cells):
                break
            current_row.cells[col_idx].text = value

    set_normal_style_in_second_column(target_table, skip_header=True)
    remove_numbering_in_second_column(target_table, skip_header=True)

    save_path = _ensure_docx_extension(output_path or dst_docx_path)
    Path(save_path).parent.mkdir(parents=True, exist_ok=True)
    dst.save(save_path)


def set_normal_style_in_second_column(table: Table, skip_header: bool = True) -> int:
    """
    Set paragraph style to 'Normal' for all paragraphs in the second column (Headline).
    Returns the number of paragraphs updated.
    """
    if not table.rows:
        return 0

    start_row = 1 if skip_header else 0
    updated = 0

    for row in table.rows[start_row:]:
        if len(row.cells) < 2:
            continue
        headline_cell = row.cells[1]
        for p in headline_cell.paragraphs:
            # Simple, robust: assign by style name
            p.style = 'Normal'
            updated += 1

    return updated


def remove_numbering_in_second_column(table, skip_header=True, non_numbered_style_name="Normal"):
    """
    Remove list numbering from the second column and prevent style-level numbering
    by assigning a non-numbered style (e.g., 'Normal').

    Returns the number of paragraphs modified.
    """
    if not table.rows:
        return 0

    start_row = 1 if skip_header else 0
    changed = 0

    for row in table.rows[start_row:]:
        if len(row.cells) < 2:
            continue

        headline_cell = row.cells[1]
        for p in headline_cell.paragraphs:
            p_elm = p._element

            # Ensure we have <w:pPr>
            pPr = p_elm.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p_elm.insert(0, pPr)

            # Remove any existing numbering and outline level
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)

            outlineLvl = pPr.find(qn('w:outlineLvl'))
            if outlineLvl is not None:
                pPr.remove(outlineLvl)

            # Assign a non-numbered style (breaks style-level numbering inheritance)
            try:
                # If the style exists, python-docx sets <w:pStyle w:val="Normal">
                p.style = p.part.document.styles[non_numbered_style_name]
            except KeyError:
                # Fallback: explicitly set <w:pStyle/> to Normal even if style not in cache
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is None:
                    pStyle = OxmlElement('w:pStyle')
                    pPr.insert(0, pStyle)
                pStyle.set(qn('w:val'), non_numbered_style_name)

            changed += 1

    return changed


# --------------------------------------------------------------------------------------------------------------------


def set_table_column_widths(
        docx_path: str,
        output_path: str = None,
        widths_cm: list[float] = None,
        expected_target_headers: list[str] = None
) -> None:
    """
    Set explicit (fixed) column widths for a target table.

    Args:
        docx_path: Path to input .docx.
        output_path: Where to save. Overwrites docx_path if None.
        widths_cm: List of column widths in centimeters. MUST match the number of columns.
        expected_target_headers: Header labels used to locate the table.

    Raises:
        ValueError: If widths_cm is missing/empty, table not found, or length mismatch.
    """
    if not widths_cm or len(widths_cm) == 0:
        raise ValueError("widths_cm must be a non-empty list of column widths (cm).")

    if expected_target_headers is None:
        expected_target_headers = ["ID"]

    doc = Document(_get_docx_path(docx_path))

    # ---- Locate target table ----
    try:
        target_table = _find_table_by_header(doc, expected_headers=expected_target_headers)
    except Exception as e:
        raise ValueError(f"Table with headers {expected_target_headers} not found: {e}")

    # ---- Validate column count ----
    if len(target_table.rows) == 0:
        raise ValueError("Target table has no rows.")
    col_count = len(target_table.rows[0].cells)
    if col_count != len(widths_cm):
        raise ValueError(f"Width count ({len(widths_cm)}) must equal column count ({col_count}).")

    # ---- Set table to fixed layout (so Word honors exact widths) ----
    tbl = target_table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Ensure <w:tblLayout w:type="fixed"/>
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')

    # Optional: set table preferred width to 'auto' by removing pct/dxa to avoid conflicts
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblPr.remove(tblW)

    # ---- Define the grid columns with exact widths ----
    def cm_to_twips(cm: float) -> int:
        return int(round((cm / 2.54) * 1440))

    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        if tblPr is not None and tblPr.getparent() is not None:
            tblPr.addnext(tblGrid)
        else:
            tbl.insert(1, tblGrid)

    # Clear existing gridCol children
    for child in list(tblGrid):
        tblGrid.remove(child)

    for cm in widths_cm:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(cm_to_twips(cm)))
        tblGrid.append(gridCol)

    # ---- Apply widths to each cell in the first row ----
    first_row = target_table.rows[0]
    for i, cm in enumerate(widths_cm):
        first_row.cells[i].width = Cm(cm)

    # Remove per-cell tcW from data rows to avoid conflicting widths
    for row in target_table.rows[1:]:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is not None:
                tcPr.remove(tcW)

    # ---- Save the document ----
    save_path = _ensure_docx_extension(output_path or docx_path)
    Path(save_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)


def set_paragraph_spacing(docx_path: str, output_path: str = None, space_before_pt: float = 0, space_after_pt: float = 3):
    """
    Set paragraph spacing ONLY for the table that follows the Section 6 heading.
    All other paragraphs and tables are left completely untouched.
    """
    doc = Document(_get_docx_path(docx_path))

    # Collect all block-level elements (paragraphs and tables) in body order
    body = doc.element.body
    children = list(body)

    # Find the index of the Section 6 heading paragraph
    section6_idx = None
    for i, child in enumerate(children):
        if child.tag.endswith('}p'):  # it's a paragraph
            text = ''.join(node.text or '' for node in child.iter() if node.tag.endswith('}t'))
            if re.search(r'^\s*6[.\)\-]?\s', text) or 'section 6' in text.lower():
                section6_idx = i
                break

    if section6_idx is None:
        return

    # Find the first table that appears after the Section 6 heading
    target_table = None
    for child in children[section6_idx + 1:]:
        if child.tag.endswith('}tbl'):
            # Wrap the raw element as a python-docx Table object
            target_table = Table(child, doc)
            break

    if target_table is None:
        return

    # Apply spacing only to paragraphs inside the target table
    for row in target_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(space_before_pt)
                paragraph.paragraph_format.space_after = Pt(space_after_pt)

    save_path = _ensure_docx_extension(output_path or docx_path)
    Path(save_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)