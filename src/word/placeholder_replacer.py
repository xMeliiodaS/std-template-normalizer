import os
from docx import Document
from docx.table import Table
from src.config.config_provider import ConfigProvider
from src.config.constants import DOCX_EXTENSION, ConfigKeys
from src.word.placeholder_values import (
    DELETE_ROW_MARKER,
    build_placeholder_replacements,
    is_report,
)
from src.config.logging_config import get_logger
import re
from pathlib import Path

logger = get_logger(__name__)


def _replace_text_in_paragraph(paragraph, replacements: dict):
    if not paragraph.runs:
        return
    # Mutate only the run text that directly contains a placeholder.
    # This avoids reflowing text between runs and preserves line breaks,
    # spacing, and run-level formatting outside the exact replacement span.
    for run in paragraph.runs:
        run_text = run.text
        new_run_text = run_text

        for placeholder, value in replacements.items():
            if placeholder in new_run_text:
                new_run_text = new_run_text.replace(placeholder, value)

        if new_run_text != run_text:
            run.text = new_run_text

    def _replace_token_across_runs(token: str, replacement: str):
        if not token:
            return

        while True:
            run_texts = [run.text for run in paragraph.runs]
            full_text = "".join(run_texts)
            start_idx = full_text.find(token)
            if start_idx == -1:
                return

            end_idx = start_idx + len(token)

            # Map absolute paragraph offsets to run index and run-local offset.
            cumulative = 0
            start_run_idx = start_off = end_run_idx = end_off = 0
            start_found = end_found = False

            for idx, run_text in enumerate(run_texts):
                next_cumulative = cumulative + len(run_text)

                if not start_found and start_idx < next_cumulative:
                    start_run_idx = idx
                    start_off = start_idx - cumulative
                    start_found = True

                if not end_found and end_idx <= next_cumulative:
                    end_run_idx = idx
                    end_off = end_idx - cumulative
                    end_found = True
                    break

                cumulative = next_cumulative

            if not (start_found and end_found):
                return

            if start_run_idx == end_run_idx:
                run = paragraph.runs[start_run_idx]
                run.text = run.text[:start_off] + replacement + run.text[end_off:]
                continue

            start_run = paragraph.runs[start_run_idx]
            end_run = paragraph.runs[end_run_idx]

            start_prefix = start_run.text[:start_off]
            end_suffix = end_run.text[end_off:]

            start_run.text = start_prefix + replacement

            for mid_idx in range(start_run_idx + 1, end_run_idx):
                paragraph.runs[mid_idx].text = ""

            end_run.text = end_suffix

    # Replace placeholders individually while preserving existing run structure.
    for placeholder, value in replacements.items():
        _replace_token_across_runs(placeholder, value)


def replace_text_in_table(table, replacements: dict):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_text_in_paragraph(paragraph, replacements)


def delete_rows_with_marker(table, marker: str):
    rows_to_delete = []

    for row in table.rows:
        row_text = " ".join(cell.text for cell in row.cells)
        if marker in row_text:
            rows_to_delete.append(row)

    for row in rows_to_delete:
        tbl = row._element.getparent()
        tbl.remove(row._element)

def _delete_paragraphs_containing_all(doc, markers: list):
    """Delete body paragraphs whose text contains ALL specified markers."""
    for paragraph in list(doc.paragraphs):
        text = paragraph.text or ""
        if all(marker in text for marker in markers):
            p_element = paragraph._element
            p_element.getparent().remove(p_element)


def replace_placeholders_using_config(docx_path, output_path=None):
    config = ConfigProvider.load_config_json()

    # Ensure output_path has .docx extension if it doesn't
    if output_path and not output_path.endswith(DOCX_EXTENSION):
        output_path = output_path + DOCX_EXTENSION

    # Ensure docx_path exists and has .docx extension
    if not docx_path.endswith(DOCX_EXTENSION):
        if os.path.exists(docx_path + DOCX_EXTENSION):
            docx_path = docx_path + DOCX_EXTENSION
        else:
            raise ValueError(f"Document path must be a {DOCX_EXTENSION} file: {docx_path}")

    out = output_path or docx_path
    logger.info("Replacing placeholders. Input: %s, Output: %s", docx_path, out)

    doc = Document(docx_path)

    report_mode = is_report(config)
    replacements = build_placeholder_replacements(config)

    # ---- Delete protocol-only cover page paragraph ----
    if not report_mode:
        _delete_paragraphs_containing_all(
            doc,
            ["ADD_PROTOCOL_NUMBER#", "ADD_STX_NUMBER"],
        )

    # ---- Body ----
    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)


    for table in doc.tables:
        if not report_mode:
            delete_rows_with_marker(
                table,
                DELETE_ROW_MARKER,
            )

        replace_text_in_table(table, replacements)

    # ---- Headers & Footers ----
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_text_in_paragraph(paragraph, replacements)

        for table in section.header.tables:
            replace_text_in_table(table, replacements)

        for paragraph in section.footer.paragraphs:
            _replace_text_in_paragraph(paragraph, replacements)

        for table in section.footer.tables:
            replace_text_in_table(table, replacements)

    for table in doc.tables:
        replace_text_in_table(
            table,
            {DELETE_ROW_MARKER: ""},
        )

    # Save document
    save_path = output_path or docx_path
    Path(save_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)
    logger.info("Output: Placeholders replaced successfully.")
