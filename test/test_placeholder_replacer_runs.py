import unittest
from docx import Document

from src.word.placeholder_replacer import _replace_text_in_paragraph


class TestPlaceholderReplacementAcrossRuns(unittest.TestCase):
    def test_replaces_placeholder_split_across_runs(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Before ")
        paragraph.add_run("ADD_")
        paragraph.add_run("DOC")
        paragraph.add_run("_TYPE#")
        paragraph.add_run(" after")

        _replace_text_in_paragraph(paragraph, {"ADD_DOC_TYPE#": "Design"})

        self.assertEqual(paragraph.text, "Before Design after")
        self.assertEqual(paragraph.runs[0].text, "Before ")
        self.assertEqual(paragraph.runs[1].text, "Design")
        self.assertEqual(paragraph.runs[2].text, "")
        self.assertEqual(paragraph.runs[3].text, "")
        self.assertEqual(paragraph.runs[4].text, " after")

    def test_replaces_placeholder_within_single_run(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("X ADD_DOC_RECORD# Y")

        _replace_text_in_paragraph(paragraph, {"ADD_DOC_RECORD#": "Protocol"})

        self.assertEqual(paragraph.text, "X Protocol Y")


if __name__ == "__main__":
    unittest.main()
